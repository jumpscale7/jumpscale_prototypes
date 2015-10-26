from docx import Document
from docx.shared import Inches

# document = Document()

document = Document('test.docx')

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote')

document.add_paragraph('first item in unordered list')

# document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in range(10):
    row_cells = table.add_row().cells
    row_cells[0].text = "test"
    row_cells[1].text = "test2"
    row_cells[2].text = str(item)

document.add_page_break()

from IPython import embed
print "DEBUG NOW ooo"
embed()


document.save('demo.docx')